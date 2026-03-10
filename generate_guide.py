from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Helpers ──
PRIMARY   = RGBColor(0x0F, 0x2C, 0x59)   # deep navy
TEAL      = RGBColor(0x00, 0xB5, 0xB5)   # accent teal
DARK_GRAY = RGBColor(0x4A, 0x55, 0x68)
BLACK     = RGBColor(0x00, 0x00, 0x00)

def add_heading(text, level=1, color=PRIMARY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(15)
        p.paragraph_format.space_before = Pt(14)
    else:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(text, bold=False, color=DARK_GRAY, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p

def add_step(number, title, detail):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    run_num = p.add_run(f"Step {number}: ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.color.rgb = TEAL
    run_title = p.add_run(title + "\n")
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = BLACK
    run_detail = p.add_run(detail)
    run_detail.font.size = Pt(10.5)
    run_detail.font.color.rgb = DARK_GRAY

def add_bullet(text, sub=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 if sub else 0.3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("📌 Note:  ")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = DARK_GRAY

def add_divider():
    p = doc.add_paragraph("─" * 70)
    p.runs[0].font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    p.runs[0].font.size = Pt(9)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

# ═══════════════════════════════════
#  COVER / TITLE
# ═══════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("BAWO Foundation")
title_run.font.size = Pt(28)
title_run.font.color.rgb = PRIMARY
title_run.bold = True

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("Website Content Management System")
sub_run.font.size = Pt(16)
sub_run.font.color.rgb = TEAL
sub_run.bold = True

sub2_p = doc.add_paragraph()
sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2_run = sub2_p.add_run("Step-by-Step User Guide")
sub2_run.font.size = Pt(13)
sub2_run.font.color.rgb = DARK_GRAY

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run("March 2026  |  Last Updated by BAWO Dev Team")
date_run.font.size = Pt(10)
date_run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

doc.add_page_break()

# ═══════════════════════════════════
#  SECTION 1 — Introduction
# ═══════════════════════════════════
add_heading("1. Introduction", level=1)
add_body(
    "The BAWO Foundation website has a built-in Content Management System (CMS) — a private "
    "online dashboard that allows authorized administrators to update text, images, and partner "
    "information across all pages of the website without any coding knowledge.",
    size=11
)
add_body(
    "The dashboard is located at:  https://bawofoundation.org/dashboard.html",
    bold=True, color=PRIMARY, size=11
)
doc.add_paragraph()

add_heading("What You Can Manage:", level=3, color=TEAL)
add_bullet("Homepage — hero title, description, pillar cards")
add_bullet("About Us — leadership bios, team photos, mission text")
add_bullet("Our Impact — project titles, descriptions, outreach images")
add_bullet("Gallery — up to 10 gallery images")
add_bullet("Partners & Giving — partner names, descriptions, logos")
add_bullet("Donate page — header text and giving options copy")
add_bullet("Footer — tagline, email, and copyright")

add_divider()

# ═══════════════════════════════════
#  SECTION 2 — Logging In
# ═══════════════════════════════════
add_heading("2. How to Log In", level=1)
add_body("Only authorized administrators with a registered email and password can access the dashboard.")

add_step("1", "Open your web browser",
         "Open Google Chrome, Safari, or any modern browser on your computer.")
add_step("2", "Go to the Dashboard URL",
         "Type the following address in your browser's address bar and press Enter:\n"
         "   https://bawofoundation.org/dashboard.html")
add_step("3", "Enter your credentials",
         "You will see a login screen with two fields:\n"
         "   • Email Address  — enter the admin email address\n"
         "   • Password  — enter your admin password\n"
         "Then click the blue 'Sign In' button.")
add_step("4", "You are now inside the Dashboard",
         "Once logged in, you will see the main dashboard with four tabs across the top:\n"
         "   Overview  |  Content  |  Donations  |  Signups")

add_note("If you see an error message saying 'Invalid credentials', double-check that your email and "
         "password are correct. Passwords are case-sensitive.")

add_divider()

# ═══════════════════════════════════
#  SECTION 3 — Dashboard Overview
# ═══════════════════════════════════
add_heading("3. Understanding the Dashboard", level=1)
add_body("After logging in, you will see a clean dashboard with the following areas:")

add_heading("Top Bar (always visible)", level=3, color=TEAL)
add_bullet("BAWO Foundation Logo — click to return to the dashboard home")
add_bullet("Website link — opens the live website in a new tab")
add_bullet("Your username badge — shows who is logged in")
add_bullet("Logout button — click this when you are done")

add_heading("Stat Cards (Overview Tab)", level=3, color=TEAL)
add_bullet("Published Content — number of content sections that have been updated")
add_bullet("Donations — count of donation activity logged")
add_bullet("Signups — number of people who signed up through the website form")
add_bullet("Comments — community feedback received")

add_heading("Tab Navigation", level=3, color=TEAL)
add_bullet("Overview — summary statistics")
add_bullet("Content — the main editor for updating site content (most important tab)")
add_bullet("Donations — view donation activity log")
add_bullet("Signups — view newsletter/volunteer sign-up entries")

add_divider()

# ═══════════════════════════════════
#  SECTION 4 — Updating Content
# ═══════════════════════════════════
add_heading("4. How to Update Website Content", level=1)
add_body("This is the most important part of the dashboard. Follow these steps carefully.")

add_step("1", "Click the 'Content' tab",
         "At the top of the dashboard, click the button labelled 'Content'. "
         "This will show you the content editor and all currently published content cards below it.")

add_step("2", "Choose the section you want to edit",
         "In the editor panel, find the dropdown menu labelled 'Site Section'.\n"
         "Click on it and you will see a list of all editable sections, organised by page:\n\n"
         "   • Home Page (hero title, descriptions, pillar cards, etc.)\n"
         "   • About Us Page (team bios, mission)\n"
         "   • Impact Page (project titles, outreach text)\n"
         "   • Gallery (images 1–10)\n"
         "   • Partners & Donate (partner info, donation copy)\n\n"
         "Select the section you want to update from this list.")

add_step("3", "The editor will auto-fill",
         "Once you select a section, the Title and Body Text fields will automatically "
         "fill in with whatever text is currently on the live website. "
         "This makes it easy to see what is there and make changes.")

add_step("4", "Make your edits",
         "   • Title field — update the heading/title for this section\n"
         "   • Body Text field — update the main paragraph text for this section\n\n"
         "You do not need to fill in both fields. You can update just the title, just "
         "the body text, or just upload a new image — whichever applies.")

add_step("5", "Upload an image (optional)",
         "If you want to update an image for this section:\n"
         "   • Click the upload area (or drag and drop an image file onto it)\n"
         "   • Select an image from your computer (JPG, PNG)\n"
         "   • The filename will appear confirming it is ready\n"
         "   You can upload an image without filling in any text fields.")

add_step("6", "Click 'Publish'",
         "When you are happy with your changes, click the teal 'Publish' button "
         "at the bottom right of the editor.\n"
         "A green toast message will appear at the bottom of the screen saying the section is now live.")

add_step("7", "Verify on the live site",
         "Click the 'Website' link in the top bar to open the live site and confirm your changes are visible.")

add_note("If you click 'Clear', it will reset the editor form but will NOT undo changes already published to the site.")

add_divider()

# ═══════════════════════════════════
#  SECTION 5 — Published Content Cards
# ═══════════════════════════════════
add_heading("5. Viewing & Re-Editing Published Content", level=1)
add_body(
    "Below the editor, you will see a grid of content cards — one for each section that has "
    "ever been published. Each card shows a preview of the content and any associated image."
)
add_bullet("Click any card to load that section back into the editor above so you can update it again.")
add_bullet("To delete a section, hover over its card and click the red trash icon that appears in the top-right corner of the card.")
add_bullet("Use the dropdown at the top of the editor to filter cards by section for easier navigation.")

add_note("Deleting a section from the dashboard removes those custom updates. "
         "The website will fall back to the original default text for that section.")

add_divider()

# ═══════════════════════════════════
#  SECTION 6 — Managing Partners
# ═══════════════════════════════════
add_heading("6. Adding & Removing Partners", level=1)
add_body("The Partners & Giving page can be managed directly from the Content tab.")

add_heading("To Add a New Partner:", level=3, color=TEAL)
add_step("1", "Go to the Content tab", "Click 'Content' in the tab navigation.")
add_step("2", "Click 'Add New Partner'", "You will find this button at the top right of the Published Content section.")
add_step("3", "Fill in the details",
         "The editor will pre-fill with placeholder text. Replace with:\n"
         "   • Title — the partner organisation's name\n"
         "   • Body Text — a short description of what they do\n"
         "   • Image upload — their logo (optional)")
add_step("4", "Click Publish", "The new partner card will appear on the live Partners page.")

add_heading("To Remove a Partner:", level=3, color=TEAL)
add_step("1", "Find the partner's card", "Scroll down in the Published Content grid to find their card.")
add_step("2", "Hover and click the trash icon", "A red circle with a trash icon appears in the top-right corner of the card when you hover over it.")
add_step("3", "Confirm deletion", "A confirmation dialog will appear. Click OK to permanently remove the partner.")

add_divider()

# ═══════════════════════════════════
#  SECTION 7 — Donations & Signups
# ═══════════════════════════════════
add_heading("7. Viewing Donations & Sign-Ups", level=1)
add_body("These tabs are view-only — they show activity from the live website.")

add_heading("Donations Tab", level=3, color=TEAL)
add_bullet("Shows a log of donation activity recorded through the website")
add_bullet("Note: Primary donation processing is handled by Donorbox. This is a secondary log.")
add_bullet("Each entry shows: donor name, email, amount (if recorded), and time")

add_heading("Signups Tab", level=3, color=TEAL)
add_bullet("Shows everyone who filled out the 'Join Our Community' sign-up form on the homepage")
add_bullet("Each entry shows: name, email, phone number, and the message they wrote")
add_bullet("This is useful for growing your mailing list and following up with interested community members")

add_divider()

# ═══════════════════════════════════
#  SECTION 8 — Logging Out
# ═══════════════════════════════════
add_heading("8. How to Log Out", level=1)
add_body("Always log out when you are finished, especially on a shared computer.")
add_step("1", "Find the Logout button",
         "Look in the top-right corner of the dashboard for the 'Logout' button.")
add_step("2", "Click Logout", "You will be returned to the login screen immediately.")
add_note("If you close the browser without logging out, your session may stay active "
         "for a short time. Always use the Logout button for security.")

add_divider()

# ═══════════════════════════════════
#  SECTION 9 — Tips & Best Practices
# ═══════════════════════════════════
add_heading("9. Tips & Best Practices", level=1)

add_bullet("Keep text concise — website paragraphs read best at 2–4 sentences.")
add_bullet("Image sizes — use images that are at least 800px wide for best quality. Keep file sizes under 5MB.")
add_bullet("Preview your changes — always click the Website link after publishing to confirm it looks correct.")
add_bullet("Avoid formatting symbols — do not paste text directly from Word with fancy characters; type it fresh or paste as plain text.")
add_bullet("Partner logos — use images with transparent backgrounds (.PNG) for the best look on the partners page.")
add_bullet("If a section looks broken — try publishing new, clean text to overwrite any corrupted content.")
add_bullet("For help — contact your web developer at the BAWO Dev Team.")

add_divider()

# ═══════════════════════════════════
#  SECTION 10 — Quick Reference
# ═══════════════════════════════════
add_heading("10. Quick Reference Card", level=1)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Task'
hdr_cells[1].text = 'Where to find it'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = PRIMARY

rows = [
    ("Log in",                     "Visit dashboard.html → enter email & password"),
    ("Update text on any page",    "Content tab → select section → edit → Publish"),
    ("Update an image",            "Content tab → select section → upload file → Publish"),
    ("Add a new partner",          "Content tab → 'Add New Partner' button → fill in → Publish"),
    ("Remove a partner",           "Content tab → hover over card → click trash icon → confirm"),
    ("View sign-up submissions",   "Signups tab"),
    ("View donation log",          "Donations tab"),
    ("Log out",                    "Top-right corner → Logout button"),
]

for task, where in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = task
    row_cells[1].text = where

doc.add_paragraph()

# ── Footer note ──
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run("BAWO Foundation — Confidential Internal Document  |  bawofoundation.org")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

# ── Save ──
desktop = os.path.expanduser("~/Desktop")
output_path = os.path.join(desktop, "BAWO_CMS_User_Guide.docx")
doc.save(output_path)
print(f"✅ Guide saved to: {output_path}")
